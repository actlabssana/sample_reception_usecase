"""
Faithful Python port of SampleReceptionController.php extraction logic.

Ports: ActlabsFormConfig (patterns), ActlabsFieldExtractor, SampleDataExtractor,
and the Pdf/Word/Spreadsheet parser pipelines.  PDF text uses `pdftotext -layout
-nopgbrk` which is exactly the controller's preferred path (PdfTextExtractor).
Only the extraction/regex behaviour is reproduced -- no Laravel/preview code.
"""
import re, subprocess, tempfile, os, csv

I = re.I; M = re.M; S = re.S

REQUIRED_FIELDS = [
    'Carrier', 'Waybill #', '# of Packages', '# of Samples', 'Priority',
    'Confirmation of Sample Receipt', 'Special Instructions/Comments',
    'Client Name', 'Client Batch #', 'Shipment #', 'Quote #, PO #, Proforma #',
    'Project', 'Company', 'Address', 'Attn', 'Phone', 'Fax', 'E-mail',
    'Additional Report To', 'Payment Method', 'Credit Card Info',
    'Reporting & Invoicing Preferences', 'Method of Sample Return',
]

# (fieldname, [ (pattern, flags), ... ])
FIELD_PATTERNS = [
    ('Carrier', [
        (r'\bCarrier\b\s*[:\-]?\s*(.+?)(?=\s{2,}|\n\s*(Waybill|#\s*of\s*Packages|#\s*of\s*Samples)\b|$)', I|M|S),
        (r'^Carrier:\s*(.+)$', I|M)]),
    ('Waybill #', [
        (r'\bWaybill\b\s*(?:#|No|Number)?\s*[:\-]?\s*(.+?)(?=\s{2,}|\n\s*(#\s*of\s*Packages|#\s*of\s*Samples|Priority)\b|$)', I|M|S),
        (r'^Waybill\s*#?\s*[:\-]?\s*(.+)$', I|M)]),
    ('# of Packages', [
        (r'(?:#|Number)\s*of\s*Packages\s*[:\-]?\s*(\d+)', I|M|S),
        (r'Packages\s*[:\-]?\s*(\d+)', I)]),
    ('# of Samples', [
        (r'(?:#|Number)\s*of\s*Samples\s*[:\-]?\s*(\d+)', I|M|S),
        (r'Samples\s*[:\-]?\s*(\d+)', I)]),
    ('Priority', [
        (r'Priority\s*[:\-]?\s*(RUSH|Normal|URGENT|EMERGENCY|STANDARD|ROUTINE)', I),
        (r'^Priority:\s*(.+)$', I|M)]),
    ('Confirmation of Sample Receipt', [
        (r'Confirmation\s+of\s+Sample\s+Receipt\s*[:\-]?\s*(Yes|No)', I),
        (r'Sample\s+Receipt\s+Confirmation\s*[:\-]?\s*(Yes|No)', I)]),
    ('Special Instructions/Comments', [
        (r'Special\s+Instructions?\s*/?\s*Comments?\s*[:\-]?\s*(.+?)(?=\n\n|\n[A-Z][a-z]+\s*:|\Z)', I|M|S),
        (r'Comments?\s*[:\-]?\s*(.+?)(?=\n\n|\n[A-Z][a-z]+\s*:|\Z)', I|M|S)]),
    ('Client Name', [
        (r'Client\s+Name\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M),
        (r'^Client:\s*(.+)$', I|M)]),
    ('Client Batch #', [
        (r'Client\s+Batch\s*#?\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M),
        (r'Batch\s*#?\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M)]),
    ('Shipment #', [
        (r'Shipment\s*#?\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M)]),
    ('Quote #, PO #, Proforma #', [
        (r'(?:Quote|PO|Proforma)\s*#?\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M)]),
    ('Project', [
        (r'^\s*Project\s*[:\-]\s*(.+)$', I|M),
        (r'Project\s+Name\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M)]),
    ('Company', [
        (r'^\s*Company\s*[:\-]\s*(.+)$', I|M),
        (r'Company\s+Name\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M)]),
    ('Address', [
        (r'^\s*Address\s*[:\-]\s*(.+)$', I|M),
        (r'Mailing\s+Address\s*[:\-]?\s*(.+?)(?=\n\n|\n[A-Z][a-z]+\s*:)', I|M|S)]),
    ('Attn', [
        (r'^\s*Attn\.?\s*[:\-]\s*(.+)$', I|M),
        (r'Attention\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M)]),
    ('Phone', [
        (r'^\s*Phone\s*[:\-]\s*([+\d][\d\s\-().]+)$', I|M),
        (r'Tel\.?\s*[:\-]?\s*([+\d][\d\s\-().]+)', I|M)]),
    ('Fax', [
        (r'^\s*Fax\s*[:\-]\s*([+\d][\d\s\-().]+)$', I|M)]),
    ('E-mail', [
        (r'^\s*E-?mail\s*[:\-]\s*([^\s].+?)$', I|M),
        (r'Email\s*[:\-]?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', I|M)]),
    ('Additional Report To', [
        (r'Additional\s+Report\s+To\s*[:\-]?\s*(.+?)(?=\n\n|\n[A-Z][a-z]+\s*:|\Z)', I|M|S)]),
    ('Payment Method', [
        (r'Payment\s+Method\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M),
        (r'Method\s+of\s+Payment\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M)]),
    ('Credit Card Info', [
        (r'Credit\s+Card\s*[:\-]?\s*(.+?)(?=\n\n|\n[A-Z][a-z]+\s*:|\Z)', I|M|S)]),
    ('Reporting & Invoicing Preferences', [
        (r'Reporting\s*[&/]?\s*Invoicing\s+Preferences?\s*[:\-]?\s*(.+?)(?=\n\n|\n[A-Z][a-z]+\s*:|\Z)', I|M|S)]),
    ('Method of Sample Return', [
        (r'Method\s+of\s+Sample\s+Return\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M),
        (r'Sample\s+Return\s*[:\-]?\s*(.+?)(?=\s{2,}|\n|$)', I|M)]),
    ('Prep Code', [
        (r'\bPrep\.?\s*Code\b\s*[:\-]?\s*(.+?)(?=\s{2,}|\n\s*(Analysis|Elements|Sample)\b|$)', I|M|S)]),
    ('Analysis Code / Elements', [
        (r'\b(Analysis\s*Code|Elements)\b\s*[:\-]?\s*(.+?)(?=\s{2,}|\n\s*(Sample|Prep)\b|$)', I|M|S)]),
]

PAYMENT_PATTERNS = [
    ('Payment included', r'Payment\s+is\s+included'),
    ('New Credit Card', r'Charge\s+to\s+NEW\s+Credit\s+Card'),
    ('Credit Card on file', r'Charge\s+to\s+Credit\s+Card\s+on\s+file'),
    ('Established Credit', r'Credit\s+has\s+been\s+established'),
]

SAMPLE_HEADER_MAPPINGS = [
    (r'Sample\s+Numbers', 'Sample Numbers'),
    (r'Sample\s+Type', 'Sample Type'),
    (r'Prep\.?\s*Code', 'Prep Code'),
    (r'Analysis\s+Code|Elements', 'Analysis Code / Elements'),
    (r'No\.?\s*of\s*Samples|#\s*Samples|Qty', '# of Samples'),
]

SKIP_PATTERNS = [
    r'^Page', r'^Rev\.', r'Activation\s+Laboratories', r'Authorized\s+Signature',
    r'Client\s+Name', r'Sample\s+Numbers|Sample\s+Type|Prep\.?\s*Code|Analysis\s+Code|Elements',
]

COMMON_MAP = [
    (r'^Carrier$', 'Carrier'),
    (r'^Waybill\s*#?$', 'Waybill #'),
    (r'^(?:#|Number)\s*of\s*Packages$', '# of Packages'),
    (r'^(?:#|Number)\s*of\s*Samples$', '# of Samples'),
    (r'^Priority$', 'Priority'),
    (r'^Confirmation\s+of\s+Sample\s+Receipt$', 'Confirmation of Sample Receipt'),
    (r'^Sample\s+Receipt\s+Confirmation$', 'Confirmation of Sample Receipt'),
    (r'^Special\s+Instructions?$', 'Special Instructions/Comments'),
    (r'^Comments?$', 'Special Instructions/Comments'),
    (r'^Client\s+Name$', 'Client Name'),
    (r'^Client$', 'Client Name'),
    (r'^Client\s+Batch\s*#?$', 'Client Batch #'),
    (r'^Batch\s*#?$', 'Client Batch #'),
    (r'^Shipment\s*#?$', 'Shipment #'),
    (r'^Quote\s*#?.*$', 'Quote #, PO #, Proforma #'),
    (r'^PO\s*#?$', 'Quote #, PO #, Proforma #'),
    (r'^Proforma\s*#?$', 'Quote #, PO #, Proforma #'),
    (r'^Project\s*(Name)?$', 'Project'),
    (r'^Company\s*(Name)?$', 'Company'),
    (r'^Address$', 'Address'),
    (r'^Mailing\s+Address$', 'Address'),
    (r'^Attn\.?$', 'Attn'),
    (r'^Attention$', 'Attn'),
    (r'^Phone$', 'Phone'),
    (r'^Tel\.?$', 'Phone'),
    (r'^Telephone$', 'Phone'),
    (r'^Fax$', 'Fax'),
    (r'^E-?mail$', 'E-mail'),
    (r'^Email$', 'E-mail'),
    (r'^Method\s+of\s+Payment$', 'Payment Method'),
    (r'^Payment\s+Method$', 'Payment Method'),
    (r'^Method\s+of\s+Sample\s+Return$', 'Method of Sample Return'),
    (r'^Sample\s+Return$', 'Method of Sample Return'),
    (r'^Additional\s+Report\s+(to|To)$', 'Additional Report To'),
    (r'^Prep\.?\s*Code$', 'Prep Code'),
    (r'^(Analysis\s*Code|Elements)$', 'Analysis Code / Elements'),
    (r'^Credit\s+Card$', 'Credit Card Info'),
    (r'^Reporting.*Invoicing$', 'Reporting & Invoicing Preferences'),
]


def init_fields():
    return {f: '' for f in REQUIRED_FIELDS}


# ---------------- ActlabsFieldExtractor ----------------
def fe_extract(content):
    fields = {}
    for name, patterns in FIELD_PATTERNS:
        for pat, flags in patterns:
            m = re.search(pat, content, flags)
            if m:
                # PHP: trim($m[1] ?? $m[2] ?? '')
                val = m.group(1) if m.group(1) is not None else (m.group(2) if m.lastindex and m.lastindex >= 2 and m.group(2) is not None else '')
                fields[name] = (val or '').strip()
                break
    fe_priority(content, fields)
    fe_confirmation(content, fields)
    fe_payment(content, fields)
    return fields


def fe_priority(content, fields):
    if fields.get('Priority', '') != '':
        if re.search(r'RUSH|URGENT|EMERGENCY|ASAP', fields['Priority'], I):
            fields['Priority'] = 'RUSH'
        elif re.search(r'NORMAL|STANDARD|ROUTINE', fields['Priority'], I):
            fields['Priority'] = 'Normal'
        return
    if re.search(r'Priority\s*[:\-]?\s*(RUSH|URGENT|EMERGENCY|ASAP)', content, I):
        fields['Priority'] = 'RUSH'
    elif re.search(r'Priority\s*[:\-]?\s*(NORMAL|STANDARD|ROUTINE)', content, I):
        fields['Priority'] = 'Normal'
    elif re.search(r'\b(RUSH|URGENT|EMERGENCY)\b', content, I):
        fields['Priority'] = 'RUSH'
    else:
        fields['Priority'] = 'Normal'


def fe_confirmation(content, fields):
    if 'Confirmation of Sample Receipt' not in fields and re.search(r'Confirmation\s+of\s+Sample\s+Receipt', content, I):
        if re.search(r'\bYes\b', content, I):
            fields['Confirmation of Sample Receipt'] = 'Yes'
        elif re.search(r'\bNo\b', content, I):
            fields['Confirmation of Sample Receipt'] = 'No'


def fe_payment(content, fields):
    for method, pat in PAYMENT_PATTERNS:
        if re.search(pat, content, I):
            fields['Payment Method'] = method
            break


def sweep_key_values(content, fields):
    pat = r'^(?P<key>[A-Za-z][A-Za-z /#&.()+\-]{2,}):\s*(?P<val>.+)$'
    for m in re.finditer(pat, content, M):
        label = m.group('key').strip()
        value = m.group('val').strip()
        if label == '' or value == '':
            continue
        extract_common_field(label, value, fields)


def _normalize_field_label(label):
    label = re.sub(r'\s+', ' ', label.strip())
    return label.lower().title()


def extract_common_field(label, value, fields):
    label = label.strip(); value = value.strip()
    if label == '' or value == '':
        return
    for pat, name in COMMON_MAP:
        if re.search(pat, label, I):
            if name not in fields or fields[name] == '' or len(value) > len(fields[name]):
                fields[name] = value
            return
    if re.search(r'^Priority$', label, I):
        if re.search(r'RUSH|URGENT|EMERGENCY|ASAP', value, I):
            fields['Priority'] = 'RUSH'
        elif re.search(r'NORMAL|STANDARD|ROUTINE', value, I):
            fields['Priority'] = 'Normal'
        else:
            fields['Priority'] = value
        return
    if re.search(r'^Confirmation\s+of\s+Sample\s+Receipt$', label, I):
        if re.search(r'\bYes\b', value, I):
            fields['Confirmation of Sample Receipt'] = 'Yes'
        elif re.search(r'\bNo\b', value, I):
            fields['Confirmation of Sample Receipt'] = 'No'
        else:
            fields['Confirmation of Sample Receipt'] = value
        return
    nl = _normalize_field_label(label)
    if nl not in fields or fields[nl] == '':
        fields[nl] = value


# ---------------- SampleDataExtractor ----------------
def is_sample_table_header(line):
    return bool(re.search(r'(Sample\s+Numbers?|Sample\s+ID|Sample\s+Type|Prep\.?\s*Code|Analysis\s*Code|Elements|Assay)', line, I))


def normalize_header(header):
    header = header.strip()
    for pat, norm in SAMPLE_HEADER_MAPPINGS:
        if re.search(pat, header, I):
            return norm
    return header


def normalize_headers(headers):
    return [normalize_header(h) for h in headers]


def _extract_headers(line):
    parts = re.split(r'\s{2,}|\t|\s\|\s', line)
    return [normalize_header(p) for p in parts if p.strip()]


def _is_non_data_row(line):
    for p in SKIP_PATTERNS:
        if re.search(p, line, I):
            return True
    return False


def _extract_row_data(line, headers):
    if line.strip() == '' or _is_non_data_row(line):
        return {}
    parts = re.split(r'\s{2,}|\t|\s\|\s', line)
    vals = [p for p in parts if p.strip()]
    if len(vals) < 2:
        return {}
    entry = {}
    for j, h in enumerate(headers):
        entry[h] = vals[j] if j < len(vals) else ''
    return entry


def sde_extract(content):
    lines = re.split(r'\r?\n', content)
    sample_data = []; inside = False; headers = []
    for line in lines:
        if not inside and is_sample_table_header(line):
            headers = _extract_headers(line); inside = True; continue
        if inside and headers:
            if is_sample_table_header(line):
                headers = _extract_headers(line); continue
            row = _extract_row_data(line, headers)
            if row:
                sample_data.append(row)
    return sample_data


def create_sample_entry(row_data, headers):
    entry = {}
    for i in range(len(headers)):
        entry[headers[i]] = row_data[i] if i < len(row_data) else ''
    return {k: v for k, v in entry.items() if v is not None and v != ''}


# ---------------- Parsers ----------------
def pdf_text(path):
    """PdfTextExtractor: pdftotext -layout -nopgbrk (preferred controller path)."""
    tmp = tempfile.mktemp(suffix='.txt')
    try:
        subprocess.run(['pdftotext', '-layout', '-nopgbrk', path, tmp],
                       capture_output=True, timeout=120)
        if os.path.isfile(tmp):
            with open(tmp, encoding='utf-8', errors='replace') as f:
                text = f.read()
            os.unlink(tmp)
            if text.strip():
                return text
    except Exception:
        pass
    return ''


def parse_pdf(path):
    text = pdf_text(path)
    fields = init_fields()
    fields.update(fe_extract(text))
    sweep_key_values(text, fields)
    sample_data = sde_extract(text)
    return {'fields': fields, 'raw': text, 'sample_data': sample_data}


def parse_word(path):
    import docx
    doc = docx.Document(path)
    text_blocks = []
    fields = init_fields()
    sample_data = []
    # Tables (mirrors WordParser.processTable)
    for table in doc.tables:
        sample_headers = []; found = False
        for row in table.rows:
            row_data = [c.text.strip() for c in row.cells]
            if not any(row_data):
                continue
            row_text = ' '.join(row_data)
            if not found and is_sample_table_header(row_text):
                sample_headers = normalize_headers(row_data)
                found = True; continue
            if found and len(row_data) >= len(sample_headers):
                entry = create_sample_entry(row_data, sample_headers)
                if entry:
                    sample_data.append(entry)
                continue
            if len(row_data) >= 2 and row_data[0]:
                extract_common_field(row_data[0], row_data[1], fields)
    # Paragraph text blocks
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            text_blocks.append(t)
    text = '\n'.join(text_blocks)
    fields.update(fe_extract(text))
    sweep_key_values(text, fields)
    if not sample_data:
        sample_data = sde_extract(text)
    return {'fields': fields, 'raw': text, 'sample_data': sample_data}


def _sheet_rows(path, ext):
    """Return list of rows; each row a list of formatted-string cell values."""
    rows = []
    if ext == 'csv':
        with open(path, encoding='utf-8', errors='replace', newline='') as f:
            for r in csv.reader(f):
                rows.append([('' if c is None else str(c)) for c in r])
        return [rows]  # single "sheet"
    if ext == 'xls':
        import xlrd
        book = xlrd.open_workbook(path)
        sheets = []
        for sh in book.sheets():
            srows = []
            for ri in range(sh.nrows):
                srows.append([_fmt(sh.cell_value(ri, ci)) for ci in range(sh.ncols)])
            sheets.append(srows)
        return sheets
    # xlsx
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    sheets = []
    for ws in wb.worksheets:
        srows = []
        for r in ws.iter_rows(values_only=True):
            srows.append([_fmt(c) for c in r])
        sheets.append(srows)
    return sheets


def _fmt(v):
    if v is None:
        return ''
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v)


def parse_spreadsheet(path, ext):
    fields = init_fields()
    sample_data = []
    text_rows = []
    for srows in _sheet_rows(path, ext):
        maxrow = len(srows)
        maxcol = max((len(r) for r in srows), default=0)

        def cell(r, c):  # 1-indexed
            if 1 <= r <= maxrow:
                row = srows[r-1]
                if 1 <= c <= len(row):
                    return row[c-1]
            return ''

        # findHeaderMap
        header_map = None
        for r in range(1, min(maxrow, 100)+1):
            hit = 0
            for c in range(1, maxcol+1):
                v = cell(r, c)
                if isinstance(v, str) and is_sample_table_header(v):
                    hit += 1
            if hit >= 2:
                cols = {}
                for c in range(1, maxcol+1):
                    v = cell(r, c)
                    if v is not None and v != '':
                        cols[c] = normalize_header(v)
                if cols:
                    header_map = {'row': r, 'cols': cols}
                    break
        if header_map:
            for r in range(header_map['row']+1, maxrow+1):
                entry = {}; has = False
                for ci, header in header_map['cols'].items():
                    val = cell(r, ci)
                    if val is not None and val != '':
                        has = True
                    entry[header] = val
                if has:
                    sample_data.append(entry)
                else:
                    break
        # text rows
        for r in range(1, maxrow+1):
            parts = [cell(r, c) for c in range(1, maxcol+1) if cell(r, c) not in (None, '')]
            if parts:
                text_rows.append(' | '.join(parts))
        # KV blocks
        for r in range(1, maxrow+1):
            for c in range(1, maxcol):
                label = cell(r, c); value = cell(r, c+1)
                if isinstance(label, str) and label != '' and value != '':
                    extract_common_field(label, value, fields)

    text = '\n'.join(text_rows)
    fields.update(fe_extract(text))
    sweep_key_values(text, fields)
    if not sample_data:
        sample_data = sde_extract(text)
    return {'fields': fields, 'raw': text, 'sample_data': sample_data}


def parse_file(path):
    ext = os.path.splitext(path)[1].lower().lstrip('.')
    if ext == 'pdf':
        return parse_pdf(path)
    if ext in ('doc', 'docx'):
        return parse_word(path)
    if ext in ('xlsx', 'xls', 'csv'):
        return parse_spreadsheet(path, ext)
    raise ValueError('Unsupported: ' + ext)
